import sys, json, requests
from PyQt5.QtWidgets import *
from PyQt5.QtCore import *

from matplotlib.backends.backend_qt5agg import FigureCanvasQTAgg as FigureCanvas
from matplotlib.figure import Figure
import matplotlib

from docx import Document

# ========= 字体修复（关键）=========
matplotlib.rcParams['font.family'] = 'Arial Unicode MS'  # Mac最稳
matplotlib.rcParams['axes.unicode_minus'] = False

# ========= 科目 =========
SUBJECTS = ["语文", "数学", "英语", "物理", "化学", "生物"]

# ========= 数据 =========
def load_data(file):
    try:
        return json.load(open(file, "r"))
    except:
        return {}

def save_data(file, data):
    json.dump(data, open(file, "w"))

users = load_data("users.json")
students = load_data("students.json")

# ========= AI =========
API_KEY = "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"

def call_ai(prompt):
    url = "https://openrouter.ai/api/v1/chat/completions"
    headers = {"Authorization": f"Bearer {API_KEY}"}

    data = {
        "model": "openai/gpt-4o-mini",
        "messages": [{"role": "user", "content": prompt}]
    }

    try:
        r = requests.post(url, json=data, headers=headers, timeout=20)
        if r.status_code != 200:
            return f"❌ AI错误\n{r.text}"
        return r.json()["choices"][0]["message"]["content"]
    except Exception as e:
        return f"❌ 网络异常：{e}"

# ========= 登录 =========
class Login(QWidget):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("AI学情分析系统")
        self.resize(350, 250)

        layout = QVBoxLayout()

        title = QLabel("🎓 AI学情分析系统")
        title.setAlignment(Qt.AlignCenter)
        title.setStyleSheet("font-size:20px;font-weight:bold;")

        self.u = QLineEdit()
        self.u.setPlaceholderText("👤 请输入用户名")

        self.p = QLineEdit()
        self.p.setPlaceholderText("🔒 请输入密码")
        self.p.setEchoMode(QLineEdit.Password)

        self.msg = QLabel("")

        btn1 = QPushButton("登录")
        btn2 = QPushButton("注册")

        btn1.clicked.connect(self.login)
        btn2.clicked.connect(self.reg)

        layout.addWidget(title)
        layout.addWidget(self.u)
        layout.addWidget(self.p)
        layout.addWidget(btn1)
        layout.addWidget(btn2)
        layout.addWidget(self.msg)

        self.setLayout(layout)

    def login(self):
        if self.u.text() in users and users[self.u.text()] == self.p.text():
            self.main = Main()
            self.main.show()
            self.close()
        else:
            self.msg.setText("❌ 登录失败")

    def reg(self):
        users[self.u.text()] = self.p.text()
        save_data("users.json", users)
        self.msg.setText("✅ 注册成功")

# ========= 主界面 =========
class Main(QWidget):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("AI学情分析系统 Pro Max")
        self.resize(1100, 650)

        main_layout = QHBoxLayout()
        menu = QVBoxLayout()
        self.stack = QStackedLayout()

        pages = ["📊 学情分析", "🧑 学生管理", "📈 成绩图表", "📄 导出报告"]

        for i, name in enumerate(pages):
            btn = QPushButton(name)
            btn.clicked.connect(lambda _, x=i: self.stack.setCurrentIndex(x))
            menu.addWidget(btn)

        menu.addStretch()

        self.stack.addWidget(self.page_analysis())
        self.stack.addWidget(self.page_students())
        self.stack.addWidget(self.page_chart())
        self.stack.addWidget(self.page_export())

        main_layout.addLayout(menu, 1)
        main_layout.addLayout(self.stack, 4)
        self.setLayout(main_layout)

    # ===== 学情分析 =====
    def page_analysis(self):
        w = QWidget()
        layout = QVBoxLayout()

        title = QLabel("📊 AI学情分析")
        title.setAlignment(Qt.AlignCenter)
        title.setStyleSheet("font-size:20px;font-weight:bold;")

        self.name = QLineEdit()
        self.name.setPlaceholderText("👤 输入学生姓名（用于保存记录）")

        layout.addWidget(title)
        layout.addWidget(self.name)

        self.inputs = {}
        for sub in SUBJECTS:
            line = QLineEdit()
            line.setPlaceholderText(f"{sub}成绩（0-100）")
            self.inputs[sub] = line
            layout.addWidget(line)

        self.output = QTextEdit()
        btn = QPushButton("🚀 生成AI分析报告")
        btn.clicked.connect(self.analyze)

        layout.addWidget(btn)
        layout.addWidget(self.output)

        w.setLayout(layout)
        return w

    def analyze(self):
        name = self.name.text()
        if not name:
            self.output.setText("❌ 请输入学生姓名")
            return

        scores = {}
        try:
            for sub, line in self.inputs.items():
                val = int(line.text())
                if not (0 <= val <= 100):
                    raise ValueError
                scores[sub] = val
        except:
            self.output.setText("❌ 成绩必须为0-100的整数")
            return

        students[name] = scores
        save_data("students.json", students)

        avg = sum(scores.values()) / len(scores)

        prompt = f"""
学生成绩：{scores}
平均分：{avg}

请输出：
1.总体评价
2.优势学科
3.薄弱学科
4.偏科情况
5.提升建议
"""

        self.output.setText("⏳ AI分析中...")
        QApplication.processEvents()

        result = call_ai(prompt)
        self.output.setText(f"📊 平均分：{avg:.2f}\n\n{result}")

        self.refresh()

    # ===== 学生管理 =====
    def page_students(self):
        w = QWidget()
        layout = QVBoxLayout()

        layout.addWidget(QLabel("📋 学生数据"))

        self.list = QListWidget()
        layout.addWidget(self.list)

        btn_clear = QPushButton("🧹 清空所有数据")
        btn_clear.clicked.connect(self.confirm_clear)
        layout.addWidget(btn_clear)

        w.setLayout(layout)
        self.refresh()
        return w

    def confirm_clear(self):
        reply = QMessageBox.question(
            self,
            "确认删除",
            "⚠️ 确定清空所有学生数据？",
            QMessageBox.Yes | QMessageBox.No
        )
        if reply == QMessageBox.Yes:
            self.clear_data()

    def clear_data(self):
        global students
        students = {}
        save_data("students.json", students)
        self.refresh()
        QMessageBox.information(self, "成功", "✅ 数据已清空")

    def refresh(self):
        if hasattr(self, "list"):
            self.list.clear()
            for k, v in students.items():
                self.list.addItem(f"{k}：{v}")

    # ===== 图表 =====
    def page_chart(self):
        w = QWidget()
        layout = QVBoxLayout()

        self.canvas = FigureCanvas(Figure(figsize=(5, 3)))
        layout.addWidget(self.canvas)

        btn = QPushButton("📈 显示最新学生成绩图")
        btn.clicked.connect(self.draw_chart)
        layout.addWidget(btn)

        w.setLayout(layout)
        return w

    def draw_chart(self):
        if not students:
            return

        name = list(students.keys())[-1]
        s = students[name]

        self.canvas.figure.clear()
        ax = self.canvas.figure.add_subplot(111)

        ax.bar(list(s.keys()), list(s.values()))
        ax.set_title(f"{name}成绩分析")

        for i, v in enumerate(s.values()):
            ax.text(i, v + 1, str(v), ha='center')

        self.canvas.draw()

    # ===== 导出 =====
    def page_export(self):
        w = QWidget()
        layout = QVBoxLayout()

        btn = QPushButton("📄 导出Word报告")
        btn.clicked.connect(self.export)
        layout.addWidget(btn)

        w.setLayout(layout)
        return w

    def export(self):
        path, _ = QFileDialog.getSaveFileName(
            self, "保存报告", "学情报告.docx", "Word文件 (*.docx)"
        )
        if not path:
            return

        doc = Document()
        doc.add_heading("📊 学情分析报告", 0)

        for name, scores in students.items():
            doc.add_heading(name, 1)
            avg = sum(scores.values()) / len(scores)
            doc.add_paragraph(f"平均分：{avg:.2f}")

            for k, v in scores.items():
                doc.add_paragraph(f"{k}：{v}")

        doc.save(path)
        QMessageBox.information(self, "成功", f"✅ 已导出到：\n{path}")

# ========= 启动 =========
if __name__ == "__main__":
    app = QApplication(sys.argv)

    app.setStyleSheet("""
    QWidget { background:#0f172a; color:#e2e8f0; }
    QLineEdit, QTextEdit {
        background:#1e293b;
        border:1px solid #334155;
        border-radius:8px;
        padding:6px;
    }
    QPushButton {
        background:#1e293b;
        padding:8px;
        border-radius:8px;
    }
    QPushButton:hover { background:#334155; }
    """)

    win = Login()
    win.show()

    sys.exit(app.exec_())